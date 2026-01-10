import os
import shutil
import uuid
import subprocess
from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import Optional, List
from dotenv import load_dotenv

load_dotenv()  # Load environment variables from .env

from langchain.chat_models import AzureChatOpenAI
from langchain_openai import AzureOpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.summarize import load_summarize_chain
from langchain.document_loaders import PyPDFLoader  # For PDFs; for DOCX, see below
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from langchain_community.utilities import SerpAPIWrapper
from langchain.schema import AIMessage, HumanMessage, SystemMessage

app = FastAPI(title="AI Research Assistant Backend", version="1.0.0")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load API keys and parameters
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
AZURE_DEPLOYMENT = os.getenv("AZURE_DEPLOYMENT")
AZURE_API_VERSION = os.getenv("AZURE_API_VERSION")
AZURE_ENDPOINT = os.getenv("AZURE_ENDPOINT")
SERPAPI_API_KEY = os.getenv("SERP_API_KEY")


if not (OPENAI_API_KEY and AZURE_DEPLOYMENT and AZURE_API_VERSION and AZURE_ENDPOINT):
    raise RuntimeError("Azure OpenAI configuration is incomplete. Please set the required environment variables.")

embedding_model = AzureOpenAIEmbeddings(
    azure_deployment="text-embedding-ada-002",
    openai_api_key=OPENAI_API_KEY,
    azure_endpoint=AZURE_ENDPOINT,
    chunk_size=1024
)

# In-memory session storage
class SessionData(BaseModel):
    history: List[dict] = []
    doc_summary: Optional[str] = None
    doc_vectors: Optional[List[List[float]]] = None
    doc_chunks: Optional[List[str]] = None

sessions = {}

# Persona settings
PERSONA_SETTINGS = {
    "factual": 0.0,
    "humorous": 0.7,
    "friendly": 0.5
}

class AskRequest(BaseModel):
    question: str
    research: bool = False

class PersonalityRequest(BaseModel):
    personality: str

@app.post("/set_personality")
def set_personality(req: PersonalityRequest):
    global sessions
    session_id = str(uuid.uuid4())
    sessions[session_id] = SessionData(history=[SystemMessage(content=f"You are a {req.personality} assistant.").dict()])
    return {"status": "ok", "session_id": session_id, "personality": req.personality.lower()}

@app.post("/upload")
async def upload_document(file: UploadFile = File(...), session_id: Optional[str] = None):
    global sessions
    if not session_id or session_id == "undefined" or session_id not in sessions:
        session_id = str(uuid.uuid4())
        sessions[session_id] = SessionData(history=[])
    session = sessions[session_id]

    filename = file.filename or "document"
    ext = os.path.splitext(filename)[1].lower()
    if ext not in [".pdf", ".docx", ".doc"]:
        raise HTTPException(status_code=400, detail="Unsupported file type. Only PDF, DOCX, or DOC files are allowed.")

    try:
        content_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read file: {e}")

    text_content = ""
    try:
        if ext == ".pdf":
            import fitz  # PyMuPDF
            pdf = fitz.open(stream=content_bytes, filetype="pdf")
            for page in pdf:
                text_content += page.get_text()
        elif ext == ".docx":
            from docx import Document
            import io
            doc = Document(io.BytesIO(content_bytes))
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        elif ext == ".doc":
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name
            try:
                result = subprocess.run(["antiword", tmp_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
                text_content = result.stdout.decode('utf-8', errors='ignore')
            except subprocess.CalledProcessError:
                raise HTTPException(status_code=500, detail="Failed to extract text from .doc file. Ensure antiword is installed.")
            finally:
                os.unlink(tmp_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting text: {e}")

    if not text_content.strip():
        raise HTTPException(status_code=400, detail="Document text is empty or could not be extracted.")

    try:
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunks = splitter.split_text(text_content)
        docs = [LangchainDocument(page_content=chunk) for chunk in chunks]
        summarize_llm = AzureChatOpenAI(
            azure_deployment=AZURE_DEPLOYMENT,
            openai_api_key=OPENAI_API_KEY,
            openai_api_type="azure",
            openai_api_version=AZURE_API_VERSION,
            azure_endpoint=AZURE_ENDPOINT,
            temperature=0.0,
            max_tokens=500,
            verbose=False,
        )
        summarize_chain = load_summarize_chain(summarize_llm, chain_type="map_reduce")
        summary_text = summarize_chain.run(docs)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Summarization failed: {e}")

    session.doc_summary = summary_text
    try:
        session.doc_chunks = chunks
        session.doc_vectors = embedding_model.embed_documents(chunks)
    except Exception as e:
        session.doc_vectors = None
        session.doc_chunks = None
        print(f"Warning: document embeddings failed: {e}")

    session.history.append({"role": "assistant", "content": summary_text})
    return {"session_id": session_id, "summary": summary_text}

@app.post("/chat")
async def chat(
    question: str = Body(..., embed=True),
    session_id: Optional[str] = Body(None, embed=True),
    personality: Optional[str] = Body("factual", embed=True),
    research: Optional[bool] = Body(False, embed=True)
):
    global sessions
    if not session_id or session_id == "undefined" or session_id not in sessions:
        session_id = str(uuid.uuid4())
        sessions[session_id] = SessionData(history=[])
    session = sessions[session_id]
    temperature = PERSONA_SETTINGS.get(personality.lower(), 0.3)
    
    # Build context from document summary if available
    doc_context = session.doc_summary if session.doc_summary else ""
    
    # Perform web research only if requested
    web_context = ""
    if research and SERPAPI_API_KEY:
        try:
            serp = SerpAPIWrapper(serpapi_api_key=SERPAPI_API_KEY, params={"engine": "bing", "gl": "us", "hl": "en"})
            search_result = serp.run(question)
        except Exception as e:
            search_result = ""
        if search_result and "I don't know" not in search_result and "No good search result" not in search_result:
            web_context = search_result

    # Build system prompt with context
    system_prompt = "You are an AI assistant."
    if doc_context:
        system_prompt += f"\nRelevant document information:\n{doc_context}"
    if web_context:
        system_prompt += f"\nWeb research information:\n{web_context}"
    
    # Build messages as dictionaries
    messages_dict = [{"role": "system", "content": system_prompt}]
    for msg in session.history:
        messages_dict.append(msg)
    messages_dict.append({"role": "user", "content": question})
    
    # Convert messages_dict to proper BaseMessage objects
    converted_messages = []
    for m in messages_dict:
        if m["role"] == "system":
            converted_messages.append(SystemMessage(content=m["content"]))
        elif m["role"] == "user":
            converted_messages.append(HumanMessage(content=m["content"]))
        elif m["role"] == "assistant":
            converted_messages.append(AIMessage(content=m["content"]))
        else:
            converted_messages.append(HumanMessage(content=m["content"]))
    
    try:
        chat_model = AzureChatOpenAI(
            azure_deployment=AZURE_DEPLOYMENT,
            openai_api_key=OPENAI_API_KEY,
            openai_api_type="azure",
            openai_api_version=AZURE_API_VERSION,
            azure_endpoint=AZURE_ENDPOINT,
            temperature=temperature,
            verbose=False,
        )
        assistant_response = chat_model(converted_messages)
        answer = assistant_response.content.strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM error: {e}")
    
    session.history.append({"role": "user", "content": question})
    session.history.append({"role": "assistant", "content": answer})
    return {"session_id": session_id, "answer": answer}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
